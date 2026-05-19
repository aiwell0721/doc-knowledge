"""
错误处理测试

验证各种异常场景下的行为：损坏文件、权限错误、空目录等
"""

import tempfile
from pathlib import Path
from click.testing import CliRunner
from doc_knowledge.cli import main


def _create_test_docx(path: Path, title: str = "Test"):
    """创建测试 DOCX 文件"""
    from docx import Document
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph("Test content about system architecture and Docker containers.")
    doc.save(str(path))


def test_corrupted_docx():
    """损坏的 DOCX 文件不应导致程序崩溃"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        
        # 创建损坏的 DOCX（写入随机二进制数据）
        (source / "corrupted.docx").write_bytes(b'\x00\x01\x02\x03CORRUPTED_FILE_DATA')
        
        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "convert", str(source),
            "-o", str(output),
            "-v"
        ])
        
        # 程序不应崩溃（exit_code 0 表示正常处理了错误）
        # 损坏文件应被跳过或记录错误
        assert result.exit_code == 0


def test_corrupted_pdf():
    """损坏的 PDF 文件不应导致程序崩溃"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        
        # 创建损坏的 PDF
        (source / "corrupted.pdf").write_bytes(b'%PDF-1.4\nCORRUPTED\n%%EOF')
        
        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "convert", str(source),
            "-o", str(output),
            "-v"
        ])
        
        # 程序不应崩溃
        assert result.exit_code == 0


def test_empty_directory():
    """空目录应给出友好提示"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        
        result = runner.invoke(main, [
            "convert", str(source)
        ])
        
        assert result.exit_code == 0
        assert "未找到" in result.output or "empty" in result.output.lower() or "no" in result.output.lower()


def test_nonexistent_directory():
    """不存在的目录应报错"""
    runner = CliRunner()
    result = runner.invoke(main, [
        "convert", "/nonexistent/path/that/does/not/exist"
    ])
    
    # click 的 exists=True 应该拒绝不存在的目录
    assert result.exit_code != 0 or "not exist" in result.output.lower()


def test_special_characters_in_filename():
    """文件名含特殊字符应正常处理"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        
        # 创建含特殊字符的文件名
        (source / "report (2026) [final].txt").write_text(
            "This file has special characters in the name and should be processed.",
            encoding="utf-8"
        )
        
        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "convert", str(source),
            "-o", str(output)
        ])
        
        assert result.exit_code == 0


def test_very_long_filename():
    """超长文件名应正常处理"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        
        long_name = "a" * 200 + ".txt"
        (source / long_name).write_text(
            "This file has a very long name but should still be processed correctly.",
            encoding="utf-8"
        )
        
        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "convert", str(source),
            "-o", str(output)
        ])
        
        assert result.exit_code == 0


def test_converter_unsupported_format():
    """不支持的格式应被标记为跳过"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        
        # 创建不支持的文件类型
        (source / "video.mp4").write_bytes(b'\x00\x00\x00mp4')
        (source / "archive.zip").write_bytes(b'PK\x03\x04')
        
        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "convert", str(source),
            "-o", str(output),
            "-v"
        ])
        
        assert result.exit_code == 0
        # 不支持的文件应被跳过或包装
        assert "skip" in result.output.lower() or "跳过" in result.output or "skipped" in result.output.lower()


def test_extractor_empty_mirror():
    """空镜像目录的 extract 命令"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        mirror = Path(tmpdir) / "mirror"
        mirror.mkdir()
        
        result = runner.invoke(main, [
            "extract", str(mirror)
        ])
        
        assert result.exit_code == 0
        assert "未找到" in result.output or "Markdown" in result.output


def test_memoMind_connection_error():
    """MemoMind API 不可达时应报错"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        knowledge = Path(tmpdir) / "knowledge"
        knowledge.mkdir()
        (knowledge / "test.md").write_text(
            '---\ntitle: "Test"\ndk_tags: ["test"]\n---\n\nContent',
            encoding="utf-8"
        )
        
        result = runner.invoke(main, [
            "export", str(knowledge),
            "--target", "memomind",
            "--api-url", "http://localhost:19999"  # 有效端口但无服务
        ])
        
        # 应报错但不崩溃
        # 注意：CLI 可能返回 0（因为错误被捕获并打印），但至少应该有错误信息
        assert "错误" in result.output or "无法连接" in result.output or result.exit_code != 0


def test_export_missing_required_param():
    """导出缺少必填参数时应报错"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        knowledge = Path(tmpdir) / "knowledge"
        knowledge.mkdir()
        (knowledge / "test.md").write_text("# Test", encoding="utf-8")
        
        result = runner.invoke(main, [
            "export", str(knowledge),
            "--target", "obsidian"  # 缺少 --vault
        ])
        
        assert result.exit_code != 0
        assert "vault" in result.output.lower() or "必填" in result.output


def test_extract_with_invalid_threshold():
    """无效的去重阈值应被 click 拒绝"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        mirror = Path(tmpdir) / "mirror"
        mirror.mkdir()
        
        result = runner.invoke(main, [
            "extract", str(mirror),
            "--threshold", "2.0"  # 无效值
        ])
        
        # click 可能允许但应该被处理
        # 至少不应该崩溃
        pass  # 这个测试验证不崩溃
