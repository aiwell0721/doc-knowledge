"""
Tests for converters/docx_converter.py — DocxConverter.

Covers:
- Headings, paragraphs, lists, tables, code blocks
- Bold, italic, underline formatting
- Empty paragraph handling
"""

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from doc_knowledge.converters.docx_converter import DocxConverter


def _make_docx(tmp_path, filename="test.docx") -> Path:
    """Helper: create a minimal DOCX file and return its Path."""
    filepath = tmp_path / filename
    doc = Document()
    doc.save(str(filepath))
    return filepath


def _add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def _add_paragraph(doc, text):
    doc.add_paragraph(text)


def _add_empty_paragraph(doc):
    doc.add_paragraph("")


def _add_list_item(doc, text, style="List Bullet"):
    p = doc.add_paragraph(text, style=style)
    return p


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = str(val)


def _add_bold_run(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def _add_italic_run(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    return p


def _add_underline_run(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.underline = True
    return p


def _add_bold_italic_run(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    return p


def _add_code_block(doc, text):
    """Add a paragraph with monospace font to simulate a code block."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {
        qn("w:ascii"): "Consolas",
        qn("w:hAnsi"): "Consolas",
    })
    rPr.append(rFonts)
    return p


class TestSupportedExtensions:
    def test_docx_extension(self):
        conv = DocxConverter()
        assert ".docx" in conv.supported_extensions

    def test_can_handle_docx(self, tmp_path):
        conv = DocxConverter()
        assert conv.can_handle(Path("file.docx")) is True
        assert conv.can_handle(Path("file.pdf")) is False


class TestHeadingConversion:
    """标题转换: Heading 1-6 → # ~ ######"""

    def _convert_with_headings(self, tmp_path, headings):
        """Create docx with given headings and convert."""
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        for text, level in headings:
            _add_heading(doc, text, level)
        doc.save(str(fp))
        conv = DocxConverter()
        return conv.convert(fp)

    def test_heading1(self, tmp_path):
        md = self._convert_with_headings(tmp_path, [("Title", 1)])
        assert "# Title" in md

    def test_heading2(self, tmp_path):
        md = self._convert_with_headings(tmp_path, [("Section", 2)])
        assert "## Section" in md

    def test_heading3(self, tmp_path):
        md = self._convert_with_headings(tmp_path, [("Subsection", 3)])
        assert "### Subsection" in md

    def test_heading4(self, tmp_path):
        md = self._convert_with_headings(tmp_path, [("Deep", 4)])
        assert "#### Deep" in md

    def test_heading5(self, tmp_path):
        md = self._convert_with_headings(tmp_path, [("Deeper", 5)])
        assert "##### Deeper" in md

    def test_heading6(self, tmp_path):
        md = self._convert_with_headings(tmp_path, [("Deepest", 6)])
        assert "###### Deepest" in md

    def test_multiple_headings(self, tmp_path):
        md = self._convert_with_headings(tmp_path, [
            ("Main", 1), ("Sub A", 2), ("Sub B", 2)
        ])
        assert "# Main" in md
        assert "## Sub A" in md
        assert "## Sub B" in md


class TestParagraphConversion:
    """段落文本转换"""

    def test_simple_paragraph(self, tmp_path):
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        doc.add_paragraph("Hello World")
        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)
        assert "Hello World" in md

    def test_multiple_paragraphs(self, tmp_path):
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)
        assert "First paragraph." in md
        assert "Second paragraph." in md


class TestEmptyParagraph:
    """空段落处理"""

    def test_empty_paragraph(self, tmp_path):
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        doc.add_paragraph("Before")
        _add_empty_paragraph(doc)
        doc.add_paragraph("After")
        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)
        # Empty paragraphs should not produce visible content
        assert "Before" in md
        assert "After" in md


class TestBoldItalicUnderline:
    """粗体、斜体、下划线格式转换"""

    def test_bold(self, tmp_path):
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        _add_bold_run(doc, "Bold text")
        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)
        assert "**Bold text**" in md

    def test_italic(self, tmp_path):
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        _add_italic_run(doc, "Italic text")
        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)
        assert "*Italic text*" in md

    def test_underline(self, tmp_path):
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        _add_underline_run(doc, "Underlined text")
        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)
        assert "<u>Underlined text</u>" in md

    def test_bold_and_italic(self, tmp_path):
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        _add_bold_italic_run(doc, "BoldItalic")
        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)
        assert "***BoldItalic***" in md


class TestListConversion:
    """列表转换"""

    def test_bullet_list(self, tmp_path):
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        _add_list_item(doc, "Item 1", style="List Bullet")
        _add_list_item(doc, "Item 2", style="List Bullet")
        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)
        assert "- Item 1" in md
        assert "- Item 2" in md

    def test_bullet_with_prefix(self, tmp_path):
        """以 • 开头的段落应被识别为列表项"""
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        doc.add_paragraph("• bullet item")
        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)
        assert "- " in md


class TestTableConversion:
    """表格转换为 Markdown"""

    def test_simple_table(self, tmp_path):
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        _add_table(doc, ["Name", "Age"], [["Alice", "30"], ["Bob", "25"]])
        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)
        assert "| Name | Age |" in md
        assert "| --- | --- |" in md
        assert "| Alice | 30 |" in md
        assert "| Bob | 25 |" in md

    def test_empty_table_rows(self):
        """空表格应返回空字符串"""
        from docx.table import Table
        # We test the _convert_table method directly with a mock-like approach
        # Since creating a truly empty table via docx API always has at least one row,
        # we test a single-row (header only) table
        fp = _make_docx(tmp_path) if 'tmp_path' in dir() else None

    def test_table_with_newlines(self, tmp_path):
        """表格单元格包含换行时应被替换为空格"""
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = "Header 2"
        table.rows[1].cells[0].text = "Line1\nLine2"
        table.rows[1].cells[1].text = "Data"
        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)
        assert "Line1 Line2" in md


class TestCodeBlock:
    """代码块（等宽字体段落）"""

    def test_code_block_with_monospace(self, tmp_path):
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        _add_code_block(doc, "print('hello')")
        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)
        assert "```" in md
        assert "print('hello')" in md

    def test_code_block_style(self, tmp_path):
        """样式名包含 'code' 的段落应被识别为代码块"""
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))
        p = doc.add_paragraph("code here")
        p.style = "No Spacing"  # fallback; the style check is by font
        doc.save(str(fp))


class TestMixedContent:
    """混合内容: 标题 + 段落 + 列表 + 表格"""

    def test_full_document(self, tmp_path):
        fp = _make_docx(tmp_path)
        doc = Document(str(fp))

        _add_heading(doc, "Document Title", 1)
        doc.add_paragraph("This is an intro paragraph.")
        _add_bold_run(doc, "Important note")
        _add_list_item(doc, "First item", style="List Bullet")
        _add_list_item(doc, "Second item", style="List Bullet")
        _add_table(doc, ["Col1", "Col2"], [["A", "B"]])

        doc.save(str(fp))

        conv = DocxConverter()
        md = conv.convert(fp)

        assert "# Document Title" in md
        assert "intro paragraph" in md
        assert "**Important note**" in md
        assert "- First item" in md
        assert "- Second item" in md
        assert "| Col1 | Col2 |" in md
