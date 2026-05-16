"""端到端集成测试"""

import pytest
from pathlib import Path
from docx import Document
from doc_knowledge import find_converter, MarkdownInjector


@pytest.fixture
def test_source_dir(tmp_path):
    """创建包含各种文件类型的测试目录"""
    src = tmp_path / "source"
    src.mkdir()
    
    # 创建 DOCX 文件
    doc = Document()
    doc.add_heading("集成测试文档", level=1)
    doc.add_paragraph("这是集成测试内容。")
    
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "列1"
    table.rows[0].cells[1].text = "列2"
    table.rows[1].cells[0].text = "值1"
    table.rows[1].cells[1].text = "值2"
    
    doc.save(str(src / "test.docx"))
    
    # 创建图片文件
    (src / "photo.png").write_bytes(b"\x89PNG\r\n\x1a\n")
    
    # 创建不支持的文件
    (src / "video.mp4").write_bytes(b"\x00" * 50)
    
    # 创建子目录
    sub = src / "subdir"
    sub.mkdir()
    (sub / "nested.docx").write_text("dummy")  # 不是真正的 DOCX，用于测试错误处理
    
    return src


def test_full_convert_pipeline(test_source_dir, tmp_path):
    """测试完整的转换流程"""
    output_dir = tmp_path / "mirror"
    
    # 扫描文件
    files = list(test_source_dir.rglob("*"))
    files = [f for f in files if f.is_file()]
    
    content_map = {}
    errors = []
    
    for f in files:
        converter = find_converter(f)
        if converter:
            try:
                markdown = converter.convert(f)
                frontmatter = f"---\ntitle: \"{f.name}\"\n---\n\n"
                content_map[f] = frontmatter + markdown
            except Exception as e:
                errors.append((f, e))
    
    # 注入
    injector = MarkdownInjector()
    stats = injector.inject(test_source_dir, output_dir, content_map)
    
    # 验证
    assert stats.converted >= 1  # 至少 test.docx 被转换
    assert stats.copied >= 1     # 至少 photo.png 被复制
    assert stats.skipped >= 1    # 至少 video.mp4 被跳过
    
    # 验证输出文件存在
    assert (output_dir / "test.docx.md").exists()
    assert (output_dir / "photo.png").exists()
    assert (output_dir / "video.mp4.md").exists()
    assert (output_dir / "summary.txt").exists()
    
    # 验证 Markdown 内容
    md_content = (output_dir / "test.docx.md").read_text(encoding="utf-8")
    assert "集成测试文档" in md_content
    assert "列1" in md_content or "列2" in md_content
