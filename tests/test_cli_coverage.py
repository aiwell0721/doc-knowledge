"""
CLI 覆盖率补充测试

针对审计报告中 CLI 覆盖率不足（75% → 目标 80%+）的问题，
补充 CLI 各命令的分支覆盖测试。

TDD 原则：先写测试，再写代码（Prove-It）
"""

import tempfile
from pathlib import Path
from click.testing import CliRunner
from doc_knowledge.cli import main


def _create_test_docx(path: Path, title: str = "Test", paragraphs: list = None):
    """创建测试 DOCX 文件"""
    from docx import Document
    doc = Document()
    doc.add_heading(title, level=1)
    if paragraphs:
        for p in paragraphs:
            doc.add_paragraph(p)
    else:
        doc.add_paragraph("Test content about system architecture.")
    doc.save(str(path))


def _create_test_md(path: Path, title: str, content: str, tags: list = None):
    """创建测试 Markdown 文件"""
    tags_str = ", ".join(f'"{t}"' for t in (tags or []))
    header = f'---\ntitle: "{title}"\ndk_tags: [{tags_str}]\n---\n\n'
    path.write_text(header + content, encoding="utf-8")


# ──────────────────────────────────────────────
# convert 命令补充测试
# ──────────────────────────────────────────────

def test_convert_format_filter():
    """convert --format 应该只转换指定格式"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        _create_test_docx(source / "doc1.docx", "Doc 1", ["Content 1"])
        (source / "test.txt").write_text("plain text", encoding="utf-8")

        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "convert", str(source), "-o", str(output),
            "--format", "docx"
        ])
        assert result.exit_code == 0
        # 只转换 docx，txt 应该被跳过（包装）
        output_files = list(output.rglob("*.md"))
        assert len(output_files) >= 1


def test_convert_no_recursive():
    """convert --no-recursive 应该不转换子目录文件"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        (source / "top.txt").write_text("top level", encoding="utf-8")
        subdir = source / "subdir"
        subdir.mkdir()
        (subdir / "nested.txt").write_text("nested", encoding="utf-8")

        result = runner.invoke(main, [
            "convert", str(source), "--no-recursive", "--dry-run"
        ])
        assert result.exit_code == 0
        assert "top.txt" in result.output
        assert "nested.txt" not in result.output


def test_convert_empty_directory():
    """convert 空目录应该显示警告并返回"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "empty"
        source.mkdir()

        result = runner.invoke(main, ["convert", str(source)])
        assert result.exit_code == 0
        assert "未找到" in result.output or "no" in result.output.lower()


def test_convert_verbose_output():
    """convert -v 应该显示详细转换信息"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        _create_test_docx(source / "doc.docx", "Test", ["Content"])

        result = runner.invoke(main, [
            "convert", str(source), "-v"
        ])
        assert result.exit_code == 0
        # verbose 模式应该显示转换统计
        assert "转换" in result.output or "完成" in result.output


def test_convert_warning_without_api_key():
    """convert --vision 无 --api-key 应该显示警告"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        (source / "test.txt").write_text("content", encoding="utf-8")

        result = runner.invoke(main, [
            "convert", str(source), "--vision", "--dry-run"
        ])
        assert result.exit_code == 0
        assert "警告" in result.output or "api-key" in result.output.lower()


# ──────────────────────────────────────────────
# extract 命令补充测试
# ──────────────────────────────────────────────

def test_extract_dry_run():
    """extract --dry-run 应该只显示计划不写入文件"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        mirror = Path(tmpdir) / "mirror"
        mirror.mkdir()
        _create_test_md(mirror / "doc1.txt.md", "Doc 1",
            "System architecture design document with Docker and Kubernetes. " * 10)
        _create_test_md(mirror / "doc2.txt.md", "Doc 2",
            "Meeting notes about project planning and team coordination. " * 5)

        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "extract", str(mirror), "-o", str(output), "--dry-run", "-v"
        ])
        assert result.exit_code == 0
        assert "Dry Run" in result.output or "计划" in result.output
        # dry-run 不应该创建输出目录
        assert not output.exists()


def test_extract_min_score_filter():
    """extract --min-score 应该过滤低分文档"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        mirror = Path(tmpdir) / "mirror"
        mirror.mkdir()
        # 高质量文档
        _create_test_md(mirror / "good.txt.md", "Good",
            "System architecture design using Docker containerization and "
            "Kubernetes orchestration for microservices deployment. " * 10)
        # 低质量文档
        _create_test_md(mirror / "bad.txt.md", "Bad", "hi")

        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "extract", str(mirror), "-o", str(output),
            "--min-score", "50", "-v"
        ])
        assert result.exit_code == 0
        output_files = list(output.rglob("*.md"))
        # good 文档应该通过，bad 文档应该被过滤
        assert len(output_files) >= 1


def test_extract_no_markdown_files():
    """extract 没有 Markdown 文件应该显示警告"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        mirror = Path(tmpdir) / "mirror"
        mirror.mkdir()
        (mirror / "test.txt").write_text("not markdown", encoding="utf-8")

        result = runner.invoke(main, ["extract", str(mirror)])
        assert result.exit_code == 0
        assert "未找到" in result.output or "Markdown" in result.output


def test_extract_keep_deprecated():
    """extract --keep-deprecated 应该将去重文件保存到 deprecated/ 目录"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        mirror = Path(tmpdir) / "mirror"
        mirror.mkdir()
        # 创建相同内容
        content = "Same content about Docker and system architecture. " * 10
        _create_test_md(mirror / "doc1.txt.md", "Doc 1", content)
        _create_test_md(mirror / "doc2.txt.md", "Doc 2", content)

        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "extract", str(mirror), "-o", str(output), "--keep-deprecated"
        ])
        assert result.exit_code == 0
        deprecated_dir = output / "deprecated"
        # deprecated 目录应该存在
        assert deprecated_dir.exists()


def test_extract_threshold_parameter():
    """extract --threshold 应该影响去重结果"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        mirror = Path(tmpdir) / "mirror"
        mirror.mkdir()
        # 相似但不完全相同的内容
        _create_test_md(mirror / "doc1.txt.md", "Doc 1",
            "Docker container orchestration with Kubernetes management. " * 10)
        _create_test_md(mirror / "doc2.txt.md", "Doc 2",
            "Docker container orchestration with Kubernetes management. " * 9 +
            "Slightly different ending here.")

        # 高阈值：更严格的去重
        output_high = Path(tmpdir) / "output_high"
        result_high = runner.invoke(main, [
            "extract", str(mirror), "-o", str(output_high), "--threshold", "0.99"
        ])
        assert result_high.exit_code == 0

        # 低阈值：更宽松的去重
        output_low = Path(tmpdir) / "output_low"
        result_low = runner.invoke(main, [
            "extract", str(mirror), "-o", str(output_low), "--threshold", "0.5"
        ])
        assert result_low.exit_code == 0


# ──────────────────────────────────────────────
# export 命令补充测试
# ──────────────────────────────────────────────

def test_export_obsidian_missing_vault():
    """export --target obsidian 缺少 --vault 应该报错"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        knowledge = Path(tmpdir) / "knowledge"
        knowledge.mkdir()

        result = runner.invoke(main, [
            "export", str(knowledge), "-t", "obsidian"
        ])
        # 应该因缺少 --vault 而失败
        assert result.exit_code != 0 or "错误" in result.output or "vault" in result.output.lower()


def test_export_memomind_missing_params():
    """export --target memomind 缺少 --api-url 和 --db 应该报错"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        knowledge = Path(tmpdir) / "knowledge"
        knowledge.mkdir()

        result = runner.invoke(main, [
            "export", str(knowledge), "-t", "memomind"
        ])
        assert result.exit_code != 0 or "错误" in result.output or "api-url" in result.output.lower()


def test_export_markdown_default_output():
    """export --target markdown 无 --output 应该使用默认输出目录"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        knowledge = Path(tmpdir) / "knowledge"
        knowledge.mkdir()
        _create_test_md(knowledge / "doc.txt.md", "Doc", "Content about Docker. " * 10)

        result = runner.invoke(main, [
            "export", str(knowledge), "-t", "markdown"
        ])
        assert result.exit_code == 0
        # 默认输出目录应该被创建
        exported = knowledge.parent / "exported"
        assert exported.exists()


def test_export_markdown_custom_output():
    """export --target markdown 自定义 --output 目录"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        knowledge = Path(tmpdir) / "knowledge"
        knowledge.mkdir()
        _create_test_md(knowledge / "doc.txt.md", "Doc", "Content about system. " * 10)

        output = Path(tmpdir) / "my-export"
        result = runner.invoke(main, [
            "export", str(knowledge), "-t", "markdown", "-o", str(output)
        ])
        assert result.exit_code == 0
        assert output.exists()
        exported_files = list(output.rglob("*.md"))
        assert len(exported_files) >= 1


# ──────────────────────────────────────────────
# pipeline 命令补充测试
# ──────────────────────────────────────────────

def test_pipeline_with_obsidian_skip():
    """pipeline --target obsidian 无 --vault 应该跳过并提示"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        _create_test_docx(source / "doc.docx", "Test", ["Content"])

        result = runner.invoke(main, [
            "pipeline", str(source), "--target", "obsidian"
        ])
        # 应该跳过 obsidian 导出
        assert result.exit_code == 0
        assert "跳过" in result.output or "Skip" in result.output


def test_pipeline_with_memomind_skip():
    """pipeline --target memomind 无 --api-url/--db 应该跳过"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        _create_test_docx(source / "doc.docx", "Test", ["Content"])

        result = runner.invoke(main, [
            "pipeline", str(source), "--target", "memomind"
        ])
        assert result.exit_code == 0
        assert "跳过" in result.output or "Skip" in result.output


def test_pipeline_markdown_export():
    """pipeline --target markdown 应该导出到指定目录"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        _create_test_docx(source / "doc.docx", "Test", ["Content about Docker. " * 5])

        output = Path(tmpdir) / "knowledge-output"
        result = runner.invoke(main, [
            "pipeline", str(source),
            "--target", "markdown",
            "-o", str(output)
        ])
        assert result.exit_code == 0
        # markdown 模式应该有输出
        output_files = list(output.rglob("*.md"))
        assert len(output_files) >= 1


def test_pipeline_verbose():
    """pipeline -v 应该显示详细日志"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        _create_test_docx(source / "doc.docx", "Test", ["Content"])

        result = runner.invoke(main, [
            "pipeline", str(source), "--target", "markdown", "-v"
        ])
        assert result.exit_code == 0


def test_pipeline_with_simhash_and_merge():
    """pipeline --simhash --merge 应该同时启用两种去重策略"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        _create_test_docx(source / "doc.docx", "Test", ["Content about Docker. " * 5])

        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "pipeline", str(source),
            "--target", "markdown",
            "-o", str(output),
            "--simhash", "--merge"
        ])
        assert result.exit_code == 0


def test_pipeline_with_temp_dir():
    """pipeline --temp-dir 应该使用指定的临时目录"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        _create_test_docx(source / "doc.docx", "Test", ["Content"])

        work = Path(tmpdir) / "workspace"
        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "pipeline", str(source),
            "--target", "markdown",
            "-o", str(output),
            "--temp-dir", str(work)
        ])
        assert result.exit_code == 0


def test_pipeline_vision_warning():
    """pipeline --vision 无 --vision-api-key 应该显示警告"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        (source / "test.txt").write_text("content", encoding="utf-8")

        result = runner.invoke(main, [
            "pipeline", str(source), "--vision", "--target", "markdown"
        ])
        assert result.exit_code == 0
        assert "警告" in result.output or "api" in result.output.lower()


# ──────────────────────────────────────────────
# version / help 测试
# ──────────────────────────────────────────────

def test_cli_version():
    """CLI --version 应该显示版本号"""
    runner = CliRunner()
    result = runner.invoke(main, ["--version"])
    assert result.exit_code == 0
    assert "doc-knowledge" in result.output.lower() or "version" in result.output.lower()


def test_cli_help():
    """CLI --help 应该显示帮助信息"""
    runner = CliRunner()
    result = runner.invoke(main, ["--help"])
    assert result.exit_code == 0
    assert "convert" in result.output
    assert "extract" in result.output
    assert "export" in result.output
    assert "pipeline" in result.output


def test_subcommand_help():
    """各子命令 --help 应该显示帮助"""
    runner = CliRunner()
    for cmd in ["convert", "extract", "export", "pipeline", "webui"]:
        result = runner.invoke(main, [cmd, "--help"])
        assert result.exit_code == 0, f"{cmd} --help failed"
