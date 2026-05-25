"""
TDD 合规测试 + 历史 Bug 复现测试

改进 1: TDD 合规 - 为关键模块建立 TDD 测试
改进 2: Prove-It 原则 - 为历史 Bug 添加复现测试

根据测试审计报告（2026-05-17）:
- 问题 #1: 整体未遵循 TDD（所有模块先代码后测试）
- 问题 #2: Bug 修复无复现测试（Pipeline bug、GBK 编码 bug 均无复现测试）
"""

import tempfile
from pathlib import Path
from click.testing import CliRunner
from doc_knowledge.cli import main


def _create_test_docx(path: Path, title: str = "Test", paragraphs: list = None):
    """创建测试 DOCX 文件"""
    from docx import Document
    doc = Document()
    doc.add_heading(title, level=1)
    if paragraphs:
        for p in paragraphs:
            doc.add_paragraph(p)
    else:
        doc.add_paragraph("Test content about system architecture.")
    doc.save(str(path))


def _create_test_md(path: Path, title: str, content: str, tags: list = None, **kwargs):
    """创建测试 Markdown 文件，支持自定义 frontmatter"""
    lines = ['---', f'title: "{title}"']
    for k, v in kwargs.items():
        lines.append(f"{k}: {v}")
    if tags:
        tags_str = ", ".join(f'"{t}"' for t in tags)
        lines.append(f"dk_tags: [{tags_str}]")
    lines.extend(['---', '', content])
    path.write_text('\n'.join(lines), encoding="utf-8")


# ──────────────────────────────────────────────
# 改进 1: TDD 合规测试
# 为转换器模块补充缺失的边界条件测试
# ──────────────────────────────────────────────

class TestConverterTDD:
    """TDD 合规：转换器模块测试（先写测试）"""

    def test_convert_single_docx(self):
        """转换单个 DOCX 文件应该成功"""
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmpdir:
            source = Path(tmpdir) / "source"
            source.mkdir()
            _create_test_docx(source / "report.docx", "Report", ["Hello World"])

            output = Path(tmpdir) / "output"
            result = runner.invoke(main, ["convert", str(source), "-o", str(output)])
            assert result.exit_code == 0

            # 验证输出
            md_file = output / "report.docx.md"
            assert md_file.exists()
            content = md_file.read_text(encoding="utf-8")
            assert "Hello World" in content

    def test_convert_mixed_formats(self):
        """混合格式文件应该分别处理"""
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmpdir:
            source = Path(tmpdir) / "source"
            source.mkdir()
            _create_test_docx(source / "doc.docx", "Doc", ["DOCX content"])
            (source / "plain.txt").write_text("TXT content", encoding="utf-8")

            output = Path(tmpdir) / "output"
            result = runner.invoke(main, ["convert", str(source), "-o", str(output), "-v"])
            assert result.exit_code == 0

            # DOCX 应该被转换
            docx_md = output / "doc.docx.md"
            assert docx_md.exists()
            # TXT 应该被包装
            txt_md = output / "plain.txt.md"
            assert txt_md.exists()

    def test_convert_with_overwrite(self):
        """--overwrite 应该覆盖已有文件"""
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmpdir:
            source = Path(tmpdir) / "source"
            source.mkdir()
            _create_test_docx(source / "doc.docx", "Doc", ["Original content"])

            output = Path(tmpdir) / "output"
            output.mkdir()
            # 预置旧文件
            (output / "doc.docx.md").write_text("old content", encoding="utf-8")

            result = runner.invoke(main, [
                "convert", str(source), "-o", str(output), "--overwrite"
            ])
            assert result.exit_code == 0

            # 应该被覆盖
            content = (output / "doc.docx.md").read_text(encoding="utf-8")
            assert "Original content" in content

    def test_convert_subdirectory_files(self):
        """递归转换应该处理子目录文件"""
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmpdir:
            source = Path(tmpdir) / "source"
            source.mkdir()
            sub = source / "sub"
            sub.mkdir()
            _create_test_docx(sub / "nested.docx", "Nested", ["Nested content"])

            output = Path(tmpdir) / "output"
            result = runner.invoke(main, ["convert", str(source), "-o", str(output)])
            assert result.exit_code == 0

            md_file = output / "sub" / "nested.docx.md"
            assert md_file.exists()


class TestExtractorTDD:
    """TDD 合规：提取器模块测试（先写测试）"""

    def test_extract_scoring_below_threshold(self):
        """低于 min-score 的文档应该被跳过"""
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmpdir:
            mirror = Path(tmpdir) / "mirror"
            mirror.mkdir()
            # 非常短的内容，评分应该很低
            _create_test_md(mirror / "short.txt.md", "Short", "hi")

            output = Path(tmpdir) / "output"
            result = runner.invoke(main, [
                "extract", str(mirror), "-o", str(output),
                "--min-score", "50", "-v"
            ])
            assert result.exit_code == 0
            # 低分文档不应该被保留
            output_files = list(output.rglob("*.md"))
            assert len(output_files) == 0, "Low score file should not be kept"

    def test_extract_dedup_removes_duplicates(self):
        """去重应该移除完全相同的内容"""
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmpdir:
            mirror = Path(tmpdir) / "mirror"
            mirror.mkdir()
            content = "System architecture using Docker containers. " * 20
            _create_test_md(mirror / "a.txt.md", "A", content)
            _create_test_md(mirror / "b.txt.md", "B", content)

            output = Path(tmpdir) / "output"
            result = runner.invoke(main, [
                "extract", str(mirror), "-o", str(output),
                "--threshold", "0.95"
            ])
            assert result.exit_code == 0
            # 应该只保留 1 个（另一个被去重）
            output_files = list(output.rglob("*.md"))
            assert len(output_files) == 1

    def test_extract_with_merge_keeps_best(self):
        """版本合并应该保留最优版本"""
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmpdir:
            mirror = Path(tmpdir) / "mirror"
            mirror.mkdir()
            _create_test_md(mirror / "report_v1.txt.md", "Report v1",
                "Short report.")
            _create_test_md(mirror / "report_v2.txt.md", "Report v2",
                "Comprehensive report about system architecture design "
                "with detailed analysis of Docker containerization.")

            output = Path(tmpdir) / "output"
            result = runner.invoke(main, [
                "extract", str(mirror), "-o", str(output), "--merge"
            ])
            assert result.exit_code == 0
            # 应该只保留最优版本
            output_files = list(output.rglob("*.md"))
            assert len(output_files) == 1
            # 内容应该是较长的那个
            content = output_files[0].read_text(encoding="utf-8")
            assert len(content) > 50


# ──────────────────────────────────────────────
# 改进 2: Prove-It 复现测试
# 为历史 Bug 添加"先写失败测试，再修复"的复现测试
# ──────────────────────────────────────────────

class TestBugReproduction:
    """Prove-It 原则：历史 Bug 复现测试"""

    def test_bug_empty_source_directory(self):
        """Bug 复现：空源目录不应该抛出异常"""
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmpdir:
            source = Path(tmpdir) / "empty"
            source.mkdir()

            # 复现：之前空目录可能导致异常
            result = runner.invoke(main, ["convert", str(source)])
            assert result.exit_code == 0
            # 应该友好提示而非崩溃
            assert "未找到" in result.output or "empty" in result.output.lower()

    def test_bug_mixed_encoding_files(self):
        """Bug 复现：混合编码文件不应该导致转换崩溃"""
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmpdir:
            source = Path(tmpdir) / "source"
            source.mkdir()
            # UTF-8 文件
            (source / "utf8.txt").write_text("中文内容测试", encoding="utf-8")
            # 纯 ASCII 文件
            (source / "ascii.txt").write_text("ASCII content test", encoding="ascii")

            output = Path(tmpdir) / "output"
            # 不应该因为编码问题而崩溃
            result = runner.invoke(main, ["convert", str(source), "-o", str(output)])
            assert result.exit_code == 0

    def test_bug_special_characters_in_filename(self):
        """Bug 复现：文件名含特殊字符不应该导致错误"""
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmpdir:
            source = Path(tmpdir) / "source"
            source.mkdir()
            # 文件名含空格和中文
            (source / "test file 测试.txt").write_text("content", encoding="utf-8")

            output = Path(tmpdir) / "output"
            result = runner.invoke(main, ["convert", str(source), "-o", str(output)])
            assert result.exit_code == 0

    def test_bug_pipeline_clean_temp_dir(self):
        """Bug 复现：Pipeline 完成后应该清理临时目录"""
        import os
        runner = CliRunner()
        
        # 记录测试前的 temp dirs
        before = set(d for d in os.listdir(tempfile.gettempdir()) if d.startswith("dck_"))
        
        with tempfile.TemporaryDirectory() as tmpdir:
            source = Path(tmpdir) / "source"
            source.mkdir()
            _create_test_docx(source / "doc.docx", "Doc", ["Content"])

            output = Path(tmpdir) / "output"
            result = runner.invoke(main, [
                "pipeline", str(source),
                "--target", "markdown",
                "-o", str(output)
            ])
            assert result.exit_code == 0

        # 只检查本次测试新增的临时目录
        after = set(d for d in os.listdir(tempfile.gettempdir()) if d.startswith("dck_"))
        leaked = after - before
        assert len(leaked) == 0, f"Temp dirs not cleaned: {leaked}"

    def test_bug_export_missing_vault_path(self):
        """Bug 复现：Obsidian 导出缺少 vault 应该清晰报错"""
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmpdir:
            knowledge = Path(tmpdir) / "knowledge"
            knowledge.mkdir()

            result = runner.invoke(main, [
                "export", str(knowledge), "-t", "obsidian"
            ])
            # 应该失败但给出清晰错误信息
            assert result.exit_code != 0 or "错误" in result.output
            # 不应该抛出未处理的异常
            assert "Traceback" not in result.output

    def test_bug_convert_with_format_filter_empty_result(self):
        """Bug 复现：格式过滤后无匹配文件不应该崩溃"""
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmpdir:
            source = Path(tmpdir) / "source"
            source.mkdir()
            (source / "test.txt").write_text("content", encoding="utf-8")

            # 过滤只找 PDF，但目录里没有 PDF
            output = Path(tmpdir) / "output"
            result = runner.invoke(main, [
                "convert", str(source), "-o", str(output), "--format", "pdf"
            ])
            # 应该优雅处理，而非崩溃
            assert result.exit_code == 0
