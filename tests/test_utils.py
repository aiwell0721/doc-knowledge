"""测试工具函数"""

from pathlib import Path
from doc_knowledge.utils import make_frontmatter


def test_make_frontmatter_basic():
    """测试基本 frontmatter 生成"""
    result = make_frontmatter(
        title="test.pdf",
        source_path=Path("/docs/test.pdf"),
        original_format="pdf",
    )
    
    assert 'title: "test.pdf"' in result
    assert 'original_format: "pdf"' in result
    assert 'conversion_status: "converted"' in result
    assert 'source: "file://' in result
    assert 'test.pdf' in result
    assert result.startswith("---")


def test_make_frontmatter_skipped():
    """测试跳过状态的 frontmatter"""
    result = make_frontmatter(
        title="video.mp4",
        source_path=Path("/docs/video.mp4"),
        original_format="mp4",
        conversion_status="skipped",
        file_size="256 MB",
    )
    
    assert "⚠️" in result
    assert "暂不支持转换" in result
    assert 'file_size: "256 MB"' in result


def test_make_frontmatter_with_extra():
    """测试额外字段"""
    result = make_frontmatter(
        title="test.pdf",
        source_path=Path("/docs/test.pdf"),
        extra={"custom": "value"},
    )

    assert 'custom: "value"' in result


def test_run_convert_images_extracted_inside_frontmatter(tmp_path):
    """images_extracted 应在 YAML frontmatter 内

    回归：_run_convert 用字符串拼接把 images_extracted 追加在第二个 --- 之后，
    导致该字段落在 frontmatter 外、成为正文首行，YAML 解析器无法识别。
    """
    from click.testing import CliRunner
    from doc_knowledge.cli import main

    source = tmp_path / "src"
    source.mkdir()
    from docx import Document
    from PIL import Image, ImageDraw

    doc = Document()
    doc.add_heading("带图文档", level=1)
    img = Image.new("RGB", (200, 150), (255, 255, 255))
    d = ImageDraw.Draw(img)
    d.rectangle([20, 20, 120, 100], fill=(200, 30, 30))
    d.rectangle([40, 40, 90, 70], fill=(30, 60, 200))
    img_path = tmp_path / "chart.png"
    img.save(img_path)
    doc.add_picture(str(img_path))
    doc.save(str(source / "img.docx"))

    out = tmp_path / "out"
    result = CliRunner().invoke(main, ["convert", str(source), "-o", str(out)])
    assert result.exit_code == 0

    md = (out / "img.docx.md").read_text(encoding="utf-8")
    first = md.index("---")
    second = md.index("---", first + 3)
    frontmatter = md[:second]
    assert "images_extracted" in frontmatter
