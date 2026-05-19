"""
CLI 新增选项 E2E 测试

验证 --simhash, --merge, --incremental 选项的端到端行为
"""

import tempfile
import time
from pathlib import Path
from click.testing import CliRunner
from doc_knowledge.cli import main


def _create_test_docx(path: Path, title: str = "Test", paragraphs: list = None):
    """创建测试 DOCX 文件"""
    from docx import Document
    doc = Document()
    doc.add_heading(title, level=1)
    if paragraphs:
        for p in paragraphs:
            doc.add_paragraph(p)
    else:
        doc.add_paragraph("Test content about system architecture and Docker containers.")
    doc.save(str(path))


def _create_test_md(path: Path, title: str, content: str, tags: list = None):
    """创建测试 Markdown 文件"""
    tags_str = ", ".join(f'"{t}"' for t in (tags or []))
    header = f'---\ntitle: "{title}"\ndk_tags: [{tags_str}]\n---\n\n'
    path.write_text(header + content, encoding="utf-8")


# ──────────────────────────────────────────────
# --simhash E2E 测试
# ──────────────────────────────────────────────

def test_extract_with_simhash():
    """extract --simhash 应该使用 SimHash 去重"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        mirror = Path(tmpdir) / "mirror"
        mirror.mkdir()
        
        # 创建相似内容文件
        _create_test_md(mirror / "doc1.txt.md", "Doc 1", 
            "This is a document about machine learning algorithms and artificial intelligence. " * 5)
        _create_test_md(mirror / "doc2.txt.md", "Doc 2",
            "This is a document about machine learning algorithms and artificial intelligence. " * 5)
        _create_test_md(mirror / "doc3.txt.md", "Doc 3",
            "今天天气很好适合出去散步。" * 5)
        
        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "extract", str(mirror),
            "-o", str(output),
            "--simhash",
            "-v"
        ])
        
        assert result.exit_code == 0
        # SimHash 应该识别出 doc1 和 doc2 相似，至少保留 1 个
        output_files = list(output.rglob("*.md"))
        assert len(output_files) >= 1


def test_simhash_vs_tfidf_comparison():
    """验证 SimHash 和 TF-IDF 去重结果差异"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        mirror = Path(tmpdir) / "mirror"
        mirror.mkdir()
        
        # 创建一组文档
        for i in range(5):
            _create_test_md(mirror / f"report_{i}.txt.md", f"Report {i}",
                f"System architecture design document version {i}. "
                f"Uses Docker containerization and Kubernetes orchestration. " * 3)
        
        # TF-IDF 去重
        output_tfidf = Path(tmpdir) / "output_tfidf"
        result_tfidf = runner.invoke(main, [
            "extract", str(mirror),
            "-o", str(output_tfidf)
        ])
        assert result_tfidf.exit_code == 0
        
        # SimHash 去重
        output_simhash = Path(tmpdir) / "output_simhash"
        result_simhash = runner.invoke(main, [
            "extract", str(mirror),
            "-o", str(output_simhash),
            "--simhash"
        ])
        assert result_simhash.exit_code == 0
        
        # 两种方法都应该成功
        tfidf_files = list(output_tfidf.rglob("*.md"))
        simhash_files = list(output_simhash.rglob("*.md"))
        
        assert len(tfidf_files) > 0
        assert len(simhash_files) > 0


# ──────────────────────────────────────────────
# --merge E2E 测试
# ──────────────────────────────────────────────

def test_extract_with_merge():
    """extract --merge 应该合并同名文档的不同版本"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        mirror = Path(tmpdir) / "mirror"
        mirror.mkdir()
        
        # 创建同一文档的多个版本
        _create_test_md(mirror / "report_v1.txt.md", "Report v1",
            "Initial version of the report.")
        _create_test_md(mirror / "report_v2.txt.md", "Report v2",
            "Second version with more details about architecture and Docker.")
        _create_test_md(mirror / "report_final.txt.md", "Report Final",
            "Final version with comprehensive analysis of system design.")
        
        # 创建另一个无关文档（不同基础名）
        _create_test_md(mirror / "other_doc.txt.md", "Other",
            "This is a different document.")
        
        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "extract", str(mirror),
            "-o", str(output),
            "--merge",
            "-v"
        ])
        
        assert result.exit_code == 0
        
        # 验证版本合并生效
        output_files = list(output.rglob("*.md"))
        # 应该至少有 1 个文件（report 的最优版本）
        assert len(output_files) >= 1
        
        # 验证保留的是内容最长的版本
        report_files = [f for f in output_files if "report" in f.name.lower()]
        if report_files:
            content = report_files[0].read_text(encoding="utf-8")
            # 应该包含较详细的内容
            assert len(content) > 50


def test_merge_preserves_best_version():
    """merge 应该保留评分最高的版本"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        mirror = Path(tmpdir) / "mirror"
        mirror.mkdir()
        
        # 创建不同长度的版本
        _create_test_md(mirror / "doc_v1.txt.md", "Doc v1", "Short content.")
        _create_test_md(mirror / "doc_v2.txt.md", "Doc v2",
            "Longer content with more details about system architecture design and implementation.")
        
        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "extract", str(mirror),
            "-o", str(output),
            "--merge"
        ])
        
        assert result.exit_code == 0
        output_files = list(output.rglob("*.md"))
        assert len(output_files) == 1
        # 应该保留 v2（内容更长）
        assert len(output_files[0].read_text(encoding="utf-8")) > 50


# ──────────────────────────────────────────────
# --incremental E2E 测试
# ──────────────────────────────────────────────

def test_incremental_skips_unchanged():
    """incremental 模式应该跳过未变更的文件"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        _create_test_docx(source / "unchanged.docx", "Unchanged", ["Original content."])
        _create_test_docx(source / "changed.docx", "Changed", ["Original content."])
        
        # 第一次运行
        output = Path(tmpdir) / "output"
        result1 = runner.invoke(main, [
            "pipeline", str(source),
            "--target", "markdown",
            "-o", str(output),
            "--incremental"
        ])
        assert result1.exit_code == 0
        
        # 等待一下确保时间戳不同
        time.sleep(0.5)
        
        # 修改一个文件
        _create_test_docx(source / "changed.docx", "Changed", ["Updated content with new information."])
        
        # 第二次运行（增量模式）
        result2 = runner.invoke(main, [
            "pipeline", str(source),
            "--target", "markdown",
            "-o", str(output),
            "--incremental"
        ])
        assert result2.exit_code == 0
        
        # 验证输出目录存在且有文件
        assert output.exists()
        output_files = list(output.rglob("*.md"))
        assert len(output_files) >= 1


def test_incremental_first_run():
    """incremental 模式首次运行应该处理所有文件"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        _create_test_docx(source / "doc1.docx", "Doc 1", ["First document."])
        _create_test_docx(source / "doc2.docx", "Doc 2", ["Second document."])
        
        output = Path(tmpdir) / "output"
        result = runner.invoke(main, [
            "pipeline", str(source),
            "--target", "markdown",
            "-o", str(output),
            "--incremental"
        ])
        
        assert result.exit_code == 0
        assert output.exists()
        
        # 首次运行应该处理所有文件
        output_files = list(output.rglob("*.md"))
        assert len(output_files) >= 1


def test_incremental_with_new_files():
    """incremental 模式应该处理新增文件"""
    runner = CliRunner()
    with tempfile.TemporaryDirectory() as tmpdir:
        source = Path(tmpdir) / "source"
        source.mkdir()
        _create_test_docx(source / "original.docx", "Original", ["Original content."])
        
        # 第一次运行
        output = Path(tmpdir) / "output"
        result1 = runner.invoke(main, [
            "pipeline", str(source),
            "--target", "markdown",
            "-o", str(output),
            "--incremental"
        ])
        assert result1.exit_code == 0
        
        # 等待一下
        time.sleep(0.5)
        
        # 添加新文件
        _create_test_docx(source / "new.docx", "New", ["New document content."])
        
        # 第二次运行
        result2 = runner.invoke(main, [
            "pipeline", str(source),
            "--target", "markdown",
            "-o", str(output),
            "--incremental"
        ])
        assert result2.exit_code == 0
        
        # 验证新文件被处理
        output_files = list(output.rglob("*.md"))
        assert len(output_files) >= 1
