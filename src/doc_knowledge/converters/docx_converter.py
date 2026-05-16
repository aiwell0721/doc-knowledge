"""
DOCX 转 Markdown 转换器
"""

from pathlib import Path
from typing import ClassVar

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .base import BaseConverter


class DocxConverter(BaseConverter):
    """DOCX 转 Markdown 转换器
    
    支持转换:
    - 标题层级 (Heading 1-6 → # ~ ######)
    - 段落文本
    - 粗体、斜体、下划线
    - 有序列表、无序列表
    - 表格
    - 代码块（等宽字体段落）
    """
    
    supported_extensions: ClassVar[list[str]] = [".docx"]
    
    # 标题映射
    HEADING_MAP: ClassVar[dict[int, str]] = {
        1: "#",
        2: "##",
        3: "###",
        4: "####",
        5: "#####",
        6: "######",
    }
    
    def convert(self, filepath: Path) -> str:
        """
        将 DOCX 文件转换为 Markdown
        
        Args:
            filepath: DOCX 文件路径
            
        Returns:
            Markdown 内容字符串
        """
        doc = Document(str(filepath))
        parts: list[str] = []
        
        for element in doc.element.body:
            tag = element.tag.split("}")[-1]  # 去掉 XML namespace
            
            if tag == "p":
                # 段落
                para_xml = element
                para = None
                for p in doc.paragraphs:
                    if p._element is para_xml:
                        para = p
                        break
                
                if para is None:
                    continue
                
                text = self._process_paragraph(para)
                if text is not None:
                    parts.append(text)
                    
            elif tag == "tbl":
                # 表格
                table = None
                for t in doc.tables:
                    if t._element is element:
                        table = t
                        break
                
                if table is not None:
                    parts.append(self._convert_table(table))
        
        return "\n\n".join(parts)
    
    def _process_paragraph(self, para) -> str | None:
        """处理单个段落，返回 Markdown 文本"""
        style_name = para.style.name.lower() if para.style else ""
        
        # 标题
        if style_name.startswith("heading"):
            level = int(style_name.replace("heading", "").strip())
            if 1 <= level <= 6:
                text = self._get_run_text(para)
                return f"{self.HEADING_MAP[level]} {text}"
        
        # 空段落
        if not para.text.strip():
            return ""
        
        # 列表
        if self._is_list_item(para):
            indent = self._get_list_indent(para)
            text = self._get_run_text(para)
            return f"{indent}- {text}"
        
        # 代码块（等宽字体）
        if self._is_code_block(para):
            text = self._get_run_text(para)
            return f"```\n{text}\n```"
        
        # 普通段落
        text = self._get_run_text(para)
        return text
    
    def _get_run_text(self, para) -> str:
        """获取段落的 Markdown 格式文本"""
        parts: list[str] = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            # 应用格式
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            
            if run.underline:
                text = f"<u>{text}</u>"
            
            parts.append(text)
        
        return "".join(parts)
    
    def _is_list_item(self, para) -> bool:
        """判断是否为列表项"""
        # 检查编号
        if para._element.find(qn("w:numPr")) is not None:
            return True
        # 检查项目符号
        style_name = (para.style.name or "").lower()
        if "list" in style_name:
            return True
        # 检查特殊前缀
        text = para.text.strip()
        if text and text[0] in ("•", "‣", "⁃", "-", "*", "·"):
            return True
        return False
    
    def _get_list_indent(self, para) -> str:
        """获取列表缩进"""
        # 简化处理：根据段落缩进层级决定
        try:
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                ind = pPr.find(qn("w:ind"))
                if ind is not None:
                    left = ind.get(qn("w:left"))
                    if left:
                        level = int(left) // 720  # 约 720 twips = 1 级缩进
                        return "  " * min(level, 3)
        except (ValueError, AttributeError):
            pass
        return ""
    
    def _is_code_block(self, para) -> bool:
        """判断是否为代码块"""
        style_name = (para.style.name or "").lower()
        if "code" in style_name or "monospace" in style_name:
            return True
        # 检查所有 run 是否为等宽字体
        if para.runs:
            all_mono = all(
                self._is_monospace(run) for run in para.runs if run.text.strip()
            )
            if all_mono and len(para.runs) > 0:
                return True
        return False
    
    def _is_monospace(self, run) -> bool:
        """检查 run 是否使用等宽字体"""
        try:
            rPr = run._element.find(qn("w:rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is not None:
                    font = rFonts.get(qn("w:ascii"))
                    if font and any(
                        m in font.lower()
                        for m in ("consolas", "courier", "monospace", "menlo")
                    ):
                        return True
        except (AttributeError, TypeError):
            pass
        return False
    
    def _convert_table(self, table) -> str:
        """将表格转换为 Markdown"""
        if not table.rows:
            return ""
        
        lines: list[str] = []
        
        # 表头
        header_cells = table.rows[0].cells
        header = "| " + " | ".join(
            cell.text.replace("\n", " ").strip() for cell in header_cells
        ) + " |"
        lines.append(header)
        
        # 分隔线
        separator = "| " + " | ".join(
            "---" for _ in header_cells
        ) + " |"
        lines.append(separator)
        
        # 数据行
        for row in table.rows[1:]:
            row_text = "| " + " | ".join(
                cell.text.replace("\n", " ").strip() for cell in row.cells
            ) + " |"
            lines.append(row_text)
        
        return "\n".join(lines)
